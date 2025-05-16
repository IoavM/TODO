import streamlit as st
import qrcode
from PIL import Image
import io
import os
import time
import glob
import base64
from gtts import gTTS
from pypdf import PdfReader, PdfWriter
from io import BytesIO
import docx
import pandas as pd
import matplotlib.pyplot as plt
import wave
import numpy as np
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import tempfile
import sys
import importlib

# Verificar dependencias
def check_dependencies():
    required_packages = [
        "qrcode", "flask", "pillow", "streamlit", "pypdf", "gtts", 
        "googletrans", "elevenlabs", "python-docx", "pandas", 
        "matplotlib", "openpyxl", "reportlab", "pydub"
    ]
    
    missing = []
    
    for package in required_packages:
        try:
            importlib.import_module(package.replace("-", "_"))
        except ImportError:
            missing.append(package)
    
    if missing:
        st.error(f"Faltan las siguientes dependencias: {', '.join(missing)}")
        st.info("Instálalas con: pip install " + " ".join(missing))
        return False
    
    return True

# Configuración de la página
st.set_page_config(page_title="Aplicación Multifuncional", layout="wide")

# Crear carpeta temporal si no existe
temp_dir = "temp"
if not os.path.exists(temp_dir):
    try:
        os.makedirs(temp_dir)
    except Exception as e:
        st.error(f"Error al crear carpeta temporal: {str(e)}")
        st.info("Verifica que la aplicación tenga permisos de escritura en el directorio actual.")

# Función para limpiar archivos temporales antiguos
def remove_files(n):
    temp_files = glob.glob("temp/*")
    if len(temp_files) != 0:
        now = time.time()
        n_days = n * 86400
        for f in temp_files:
            if os.stat(f).st_mtime < now - n_days:
                os.remove(f)
                print("Deleted ", f)

# Limpiar archivos temporales más antiguos de 7 días
remove_files(7)

# Sidebar con opciones
with st.sidebar:
    st.title("Selecciona una función")
    option = st.radio(
        "Elige una herramienta:",
        ["Generador de Código QR", "Recortar PDF", "Convertidor de Texto a Voz", "Convertidor de Archivos"]
    )
    st.write(f"Opción seleccionada: {option}")  # Debug: Muestra la opción seleccionada

# 1. Generador de código QR
def qr_generator():
    st.title("Generador de Código QR")
    # Entrada del texto para generar el QR
    qr_text = st.text_input("Introduce el texto o enlace para generar el código QR")
    
    if qr_text:
        # Crear código QR
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(qr_text)
        qr.make(fit=True)
        
        # Convertir el código QR en imagen
        img = qr.make_image(fill='black', back_color='white')
        
        # Guardar la imagen en un buffer de memoria
        buffer = io.BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)  # Reiniciar el puntero del buffer al inicio
        
        # Mostrar la imagen en Streamlit
        st.image(buffer, caption="Código QR generado", use_column_width=True)
        
        # Botón de descarga
        st.download_button(
            label="Descargar código QR",
            data=buffer,
            file_name="codigo_qr.png",
            mime="image/png"
        )

# 2. Recortar PDF
def pdf_cutter():
    st.title("✂️ Recortar un PDF por páginas")
    st.write("Carga un archivo PDF y selecciona el rango de páginas que quieres extraer.")
    
    uploaded_file = st.file_uploader("📄 Cargar PDF", type="pdf")
    
    if uploaded_file:
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)
        st.info(f"Este PDF tiene {total_pages} páginas.")
        
        start_page = st.number_input("Página inicial", min_value=1, max_value=total_pages, value=1)
        end_page = st.number_input("Página final", min_value=1, max_value=total_pages, value=total_pages)
        
        if start_page > end_page:
            st.error("⚠️ La página inicial no puede ser mayor que la final.")
        else:
            if st.button("Recortar PDF"):
                writer = PdfWriter()
                for i in range(start_page - 1, end_page):
                    writer.add_page(reader.pages[i])
                
                output = BytesIO()
                writer.write(output)
                output.seek(0)
                
                st.success("✅ PDF recortado listo para descargar")
                st.download_button(
                    label="📥 Descargar PDF recortado",
                    data=output,
                    file_name="pdf_recortado.pdf",
                    mime="application/pdf"
                )

# 3. Convertidor de texto a voz
def text_to_speech_converter():
    st.title("Conversión de Texto a Audio")
    
    with st.sidebar:
        st.subheader("Escribe y/o selecciona texto para ser escuchado.")
    
    # Entrada de texto
    text = st.text_area("Ingrese el texto a escuchar.")
    
    # Selección de idioma
    option_lang = st.selectbox("Selecciona el lenguaje", ("Español", "English"))
    
    if option_lang == "Español":
        lg = 'es'
    if option_lang == "English":
        lg = 'en'
    
    tld = 'com'
    
    def text_to_speech(text, tld, lg):
        tts = gTTS(text, lang=lg)
        try:
            my_file_name = text[0:20]
        except:
            my_file_name = "audio"
        
        # Limpiar nombre de archivo
        my_file_name = ''.join(e for e in my_file_name if e.isalnum() or e == ' ')
        my_file_name = my_file_name.replace(' ', '_')
        if not my_file_name:
            my_file_name = "audio"
            
        tts.save(f"temp/{my_file_name}.mp3")
        return my_file_name, text
    
    if st.button("Convertir a Audio"):
        if text:
            result, output_text = text_to_speech(text, tld, lg)
            audio_file = open(f"temp/{result}.mp3", "rb")
            audio_bytes = audio_file.read()
            st.markdown(f"## Tu audio:")
            st.audio(audio_bytes, format="audio/mp3", start_time=0)
            
            with open(f"temp/{result}.mp3", "rb") as f:
                data = f.read()
                
            def get_binary_file_downloader_html(bin_file, file_label='File'):
                bin_str = base64.b64encode(data).decode()
                href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Download {file_label}</a>'
                return href
                
            st.markdown(get_binary_file_downloader_html(f"temp/{result}.mp3", file_label="Audio File"), unsafe_allow_html=True)
        else:
            st.warning("Por favor, ingresa algún texto para convertir a audio.")

# 4. Convertidor de archivos
def file_converter():
    try:
        st.title("🔄 Convertidor de Archivos")
        st.write("Convierte archivos entre diferentes formatos sin límites ni costos.")
        
        # Categorías de conversión
        conversion_category = st.selectbox(
            "Selecciona la categoría de conversión:",
            ["Imágenes", "Documentos", "Audio", "Hojas de cálculo"]
        )
        
        if conversion_category == "Imágenes":
            image_converter()
        elif conversion_category == "Documentos":
            document_converter()
        elif conversion_category == "Audio":
            audio_converter()
        elif conversion_category == "Hojas de cálculo":
            spreadsheet_converter()
    except Exception as e:
        st.error(f"Error en la función file_converter: {str(e)}")

# Convertidor de imágenes
def image_converter():
    st.subheader("Convertidor de Imágenes")
    
    supported_formats = ["JPG/JPEG", "PNG", "BMP", "WEBP", "TIFF", "GIF"]
    
    uploaded_file = st.file_uploader("📸 Sube una imagen", type=["jpg", "jpeg", "png", "bmp", "webp", "tiff", "gif"])
    
    if uploaded_file:
        try:
            # Mostrar la imagen original
            img = Image.open(uploaded_file)
            st.image(img, caption="Imagen original", use_column_width=True)
            
            # Información de la imagen
            st.info(f"Formato original: {img.format}, Modo: {img.mode}, Tamaño: {img.size[0]}x{img.size[1]} píxeles")
            
            # Opciones de conversión
            target_format = st.selectbox("Convertir a formato:", supported_formats)
            
            # Opciones de calidad para JPG
            quality = 90
            if target_format == "JPG/JPEG":
                quality = st.slider("Calidad (solo JPG):", 1, 100, 90)
            
            # Opciones de redimensionamiento
            resize_option = st.checkbox("Redimensionar imagen")
            new_width, new_height = img.size
            
            if resize_option:
                col1, col2 = st.columns(2)
                with col1:
                    new_width = st.number_input("Nuevo ancho (píxeles):", min_value=1, value=img.size[0])
                with col2:
                    new_height = st.number_input("Nuevo alto (píxeles):", min_value=1, value=img.size[1])
            
            # Procesamiento y conversión
            if st.button("Convertir imagen"):
                # Redimensionar si es necesario
                if resize_option:
                    img = img.resize((int(new_width), int(new_height)))
                
                # Convertir al formato objetivo
                format_map = {
                    "JPG/JPEG": "JPEG",
                    "PNG": "PNG",
                    "BMP": "BMP",
                    "WEBP": "WEBP",
                    "TIFF": "TIFF",
                    "GIF": "GIF"
                }
                
                target_format_code = format_map[target_format]
                
                # Manejar conversiones específicas
                if img.mode == "RGBA" and target_format_code == "JPEG":
                    img = img.convert("RGB")  # JPEG no soporta alpha
                
                # Guardar la imagen convertida
                output = BytesIO()
                if target_format_code == "JPEG":
                    img.save(output, format=target_format_code, quality=quality)
                else:
                    img.save(output, format=target_format_code)
                
                output.seek(0)
                
                # Extensiones de archivo para cada formato
                ext_map = {
                    "JPEG": "jpg",
                    "PNG": "png",
                    "BMP": "bmp",
                    "WEBP": "webp",
                    "TIFF": "tiff",
                    "GIF": "gif"
                }
                
                # Botón de descarga
                st.success(f"✅ Imagen convertida a {target_format}")
                st.download_button(
                    label=f"📥 Descargar imagen {target_format}",
                    data=output,
                    file_name=f"imagen_convertida.{ext_map[target_format_code]}",
                    mime=f"image/{ext_map[target_format_code]}"
                )
                
        except Exception as e:
            st.error(f"Error al procesar la imagen: {str(e)}")

# Convertidor de documentos
def document_converter():
    st.subheader("Convertidor de Documentos")
    
    st.info("Actualmente soportamos las siguientes conversiones: TXT a PDF, DOCX a TXT, PDF a TXT, DOCX a PDF, PDF a DOCX")
    
    conversion_type = st.selectbox(
        "Selecciona el tipo de conversión:",
        ["TXT a PDF", "DOCX a TXT", "PDF a TXT", "DOCX a PDF", "PDF a DOCX"]
    )
    
    if conversion_type == "TXT a PDF":
        txt_to_pdf()
    elif conversion_type == "DOCX a TXT":
        docx_to_txt()
    elif conversion_type == "PDF a TXT":
        pdf_to_txt()
    elif conversion_type == "DOCX a PDF":
        docx_to_pdf()
    elif conversion_type == "PDF a DOCX":
        pdf_to_docx()

# Funciones para conversión de documentos

def txt_to_pdf():
    uploaded_file = st.file_uploader("📄 Sube un archivo TXT", type=["txt"])
    
    if uploaded_file:
        try:
            # Leer el contenido del archivo TXT
            text_content = uploaded_file.getvalue().decode("utf-8")
            
            # Crear PDF con ReportLab
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            
            # Configurar fuente y tamaño
            c.setFont("Helvetica", 12)
            
            # Dividir el texto en líneas y escribir en el PDF
            y_position = 750  # Posición inicial Y (desde arriba)
            line_height = 14
            
            for line in text_content.split('\n'):
                # Manejar líneas largas
                words = line.split()
                current_line = ""
                
                for word in words:
                    if len(current_line + " " + word) < 80:  # Límite de caracteres por línea
                        current_line += " " + word if current_line else word
                    else:
                        c.drawString(50, y_position, current_line)
                        y_position -= line_height
                        current_line = word
                        
                        # Nueva página si es necesario
                        if y_position < 50:
                            c.showPage()
                            c.setFont("Helvetica", 12)
                            y_position = 750
                
                # Escribir la última línea de palabras
                if current_line:
                    c.drawString(50, y_position, current_line)
                    y_position -= line_height
                
                # Espacio extra entre párrafos
                y_position -= 5
                
                # Nueva página si es necesario
                if y_position < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_position = 750
            
            c.save()
            buffer.seek(0)
            
            st.success("✅ TXT convertido a PDF")
            st.download_button(
                label="📥 Descargar PDF",
                data=buffer,
                file_name="texto_convertido.pdf",
                mime="application/pdf"
            )
                
        except Exception as e:
            st.error(f"Error al convertir el archivo: {str(e)}")

def docx_to_txt():
    uploaded_file = st.file_uploader("📄 Sube un archivo DOCX", type=["docx"])
    
    if uploaded_file:
        try:
            # Crear un documento Word a partir del archivo subido
            doc = docx.Document(uploaded_file)
            
            # Extraer texto
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Mostrar vista previa
            st.text_area("Vista previa del texto extraído:", text[:1000] + ("..." if len(text) > 1000 else ""), height=300)
            
            # Botón de descarga
            if st.button("Descargar como TXT"):
                buffer = BytesIO()
                buffer.write(text.encode())
                buffer.seek(0)
                
                st.success("✅ DOCX convertido a TXT")
                st.download_button(
                    label="📥 Descargar TXT",
                    data=buffer,
                    file_name="documento_convertido.txt",
                    mime="text/plain"
                )
                
        except Exception as e:
            st.error(f"Error al convertir el archivo: {str(e)}")

def pdf_to_txt():
    uploaded_file = st.file_uploader("📄 Sube un archivo PDF", type=["pdf"])
    
    if uploaded_file:
        try:
            # Leer el PDF
            reader = PdfReader(uploaded_file)
            
            # Extraer texto
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            
            # Mostrar vista previa
            st.text_area("Vista previa del texto extraído:", text[:1000] + ("..." if len(text) > 1000 else ""), height=300)
            
            # Botón de descarga
            if st.button("Descargar como TXT"):
                buffer = BytesIO()
                buffer.write(text.encode())
                buffer.seek(0)
                
                st.success("✅ PDF convertido a TXT")
                st.download_button(
                    label="📥 Descargar TXT",
                    data=buffer,
                    file_name="pdf_convertido.txt",
                    mime="text/plain"
                )
                
        except Exception as e:
            st.error(f"Error al convertir el archivo: {str(e)}")

def docx_to_pdf():
    st.warning("⚠️ Esta función requiere LibreOffice o Microsoft Word instalado en el servidor.")
    uploaded_file = st.file_uploader("📄 Sube un archivo DOCX", type=["docx"])
    
    if uploaded_file:
        try:
            # Intentar importar docx2pdf
            try:
                from docx2pdf import convert as docx2pdf_convert
            except ImportError:
                st.error("La biblioteca docx2pdf no está instalada. Instálala con: pip install docx2pdf")
                return
                
            # Crear directorio temporal
            temp_dir = tempfile.TemporaryDirectory()
            input_path = f"{temp_dir.name}/input.docx"
            output_path = f"{temp_dir.name}/output.pdf"
            
            # Guardar el archivo subido
            with open(input_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Realizar la conversión
            if st.button("Convertir a PDF"):
                with st.spinner("⏳ Convirtiendo DOCX a PDF..."):
                    try:
                        docx2pdf_convert(input_path, output_path)
                        
                        # Leer el archivo convertido
                        with open(output_path, "rb") as f:
                            pdf_bytes = f.read()
                        
                        st.success("✅ DOCX convertido a PDF")
                        st.download_button(
                            label="📥 Descargar PDF",
                            data=pdf_bytes,
                            file_name="documento_convertido.pdf",
                            mime="application/pdf"
                        )
                        
                        # Opcional: previsualizar primera página
                        try:
                            reader = PdfReader(output_path)
                            if len(reader.pages) > 0:
                                st.write("Vista previa (primera página):")
                                st.write(reader.pages[0].extract_text()[:500] + "...")
                        except:
                            pass
                    except Exception as e:
                        st.error(f"Error durante la conversión: {str(e)}")
                        st.info("Asegúrate de tener LibreOffice o Microsoft Word instalado en el servidor.")
                    
            # Limpiar
            temp_dir.cleanup()
                
        except Exception as e:
            st.error(f"Error al convertir el archivo: {str(e)}")

def pdf_to_docx():
    uploaded_file = st.file_uploader("📄 Sube un archivo PDF", type=["pdf"])
    
    if uploaded_file:
        try:
            # Intentar importar pdf2docx
            try:
                from pdf2docx import Converter as PDF2DOCXConverter
            except ImportError:
                st.error("La biblioteca pdf2docx no está instalada. Instálala con: pip install pdf2docx")
                return
                
            # Crear directorio temporal
            temp_dir = tempfile.TemporaryDirectory()
            input_path = f"{temp_dir.name}/input.pdf"
            output_path = f"{temp_dir.name}/output.docx"
            
            # Guardar el archivo subido
            with open(input_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Opciones de conversión
            max_pages = 50  # Limite por defecto para evitar archivos muy grandes
            
            # Verificar número de páginas
            reader = PdfReader(input_path)
            total_pages = len(reader.pages)
            
            st.info(f"Este PDF tiene {total_pages} páginas.")
            
            if total_pages > max_pages:
                convert_all = st.checkbox(f"El PDF tiene más de {max_pages} páginas. ¿Convertir todas? (puede tomar tiempo)")
                pages_to_convert = total_pages if convert_all else max_pages
            else:
                pages_to_convert = total_pages
            
            # Realizar la conversión
            if st.button("Convertir a DOCX"):
                with st.spinner(f"⏳ Convirtiendo PDF a DOCX ({pages_to_convert} páginas)..."):
                    try:
                        # Configurar el convertidor
                        cv = PDF2DOCXConverter(input_path)
                        # Convertir por páginas
                        cv.convert(output_path, start=0, end=pages_to_convert)
                        # Cerrar el convertidor
                        cv.close()
                        
                        # Leer el archivo convertido
                        with open(output_path, "rb") as f:
                            docx_bytes = f.read()
                        
                        st.success("✅ PDF convertido a DOCX")
                        st.download_button(
                            label="📥 Descargar DOCX",
                            data=docx_bytes,
                            file_name="pdf_convertido.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        # Opcional: mostrar vista previa
                        try:
                            doc = docx.Document(output_path)
                            preview_text = "\n".join([p.text for p in doc.paragraphs][:20])
                            st.text_area("Vista previa del contenido:", preview_text[:500] + ("..." if len(preview_text) > 500 else ""), height=200)
                        except Exception as preview_error:
                            st.warning(f"No se pudo generar la vista previa: {str(preview_error)}")
                    except Exception as e:
                        st.error(f"Error durante la conversión: {str(e)}")
                    
            # Limpiar
            temp_dir.cleanup()
                
        except Exception as e:
            st.error(f"Error al convertir el archivo: {str(e)}")

# Convertidor de audio
def audio_converter():
    st.subheader("Convertidor de Audio")
    st.warning("⚠️ Esta función requiere FFmpeg instalado en el servidor.")
    st.write("Convierte archivos de audio entre diferentes formatos.")
    
    # Para trabajar con archivos de audio necesitaremos la biblioteca pydub
    # Muestra un mensaje si no está instalada
    try:
        from pydub import AudioSegment
        pydub_installed = True
    except ImportError:
        st.error("Se requiere la biblioteca 'pydub' para la conversión de audio. Instálala con: pip install pydub")
        st.info("También necesitarás FFmpeg instalado en tu sistema para la conversión de formatos.")
        pydub_installed = False
        return
    
    if pydub_installed:
        # Formatos soportados
        supported_formats = ["mp3", "wav", "ogg", "flac", "aac", "m4a"]
        
        uploaded_file = st.file_uploader("🎵 Sube un archivo de audio", type=supported_formats)
        
        if uploaded_file:
            # Obtener el formato original
            original_format = uploaded_file.name.split('.')[-1].lower()
            
            # Guardar temporalmente el archivo subido
            with open(f"temp/temp_audio.{original_format}", "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Cargar el audio con pydub
            try:
                audio = AudioSegment.from_file(f"temp/temp_audio.{original_format}", format=original_format)
                
                # Información del audio
                duration_seconds = len(audio) / 1000
                channels = audio.channels
                sample_width = audio.sample_width
                frame_rate = audio.frame_rate
                
                st.success("✅ Audio cargado correctamente")
                st.audio(uploaded_file, format=f"audio/{original_format}")
                
                st.info(f"Información del audio:\n" +
                       f"- Duración: {int(duration_seconds // 60)}:{int(duration_seconds % 60):02d} (min:seg)\n" +
                       f"- Canales: {channels} ({'Estéreo' if channels == 2 else 'Mono'})\n" +
                       f"- Profundidad de bits: {sample_width * 8} bits\n" +
                       f"- Frecuencia de muestreo: {frame_rate} Hz")
                
                # Opciones de conversión
                target_formats = [fmt for fmt in supported_formats if fmt != original_format]
                target_format = st.selectbox("Convertir a formato:", target_formats)
                
                # Opciones adicionales
                st.subheader("Opciones adicionales")
                
                # Ajuste de volumen
                volume_adjustment = st.slider("Ajuste de volumen (dB):", -20, 20, 0)
                
                # Recortar audio
                trim_audio = st.checkbox("Recortar audio")
                start_time, end_time = 0, duration_seconds
                
                if trim_audio:
                    col1, col2 = st.columns(2)
                    with col1:
                        start_time = st.number_input("Tiempo de inicio (segundos):", min_value=0.0, max_value=duration_seconds-1, value=0.0, step=0.1)
                    with col2:
                        end_time = st.number_input("Tiempo final (segundos):", min_value=start_time+0.1, max_value=duration_seconds, value=duration_seconds, step=0.1)
                
                if st.button("Convertir audio"):
                    try:
                        # Aplicar ajustes
                        processed_audio = audio
                        
                        # Ajustar volumen si es necesario
                        if volume_adjustment != 0:
                            processed_audio = processed_audio + volume_adjustment
                        
                        # Recortar si es necesario
                        if trim_audio:
                            processed_audio = processed_audio[int(start_time*1000):int(end_time*1000)]
                        
                        # Guardar en el nuevo formato
                        output_file = f"temp/converted_audio.{target_format}"
                        processed_audio.export(output_file, format=target_format)
                        
                        # Leer el archivo convertido
                        with open(output_file, "rb") as f:
                            converted_audio_bytes = f.read()
                        
                        st.success(f"✅ Audio convertido a {target_format.upper()}")
                        
                        # Reproducir el audio convertido
                        st.audio(converted_audio_bytes, format=f"audio/{target_format}")
                        
                        # Botón de descarga
                        st.download_button(
                            label=f"📥 Descargar audio {target_format.upper()}",
                            data=converted_audio_bytes,
                            file_name=f"audio_convertido.{target_format}",
                            mime=f"audio/{target_format}"
                        )
                    except Exception as e:
                        st.error(f"Error durante la conversión del audio: {str(e)}")
                        if "encoder" in str(e).lower() or "ffmpeg" in str(e).lower():
                            st.info("Este error puede deberse a que FFmpeg no está instalado o no está en el PATH.")
                            st.info("Instala FFmpeg siguiendo las instrucciones en: https://www.ffmpeg.org/download.html")
                    
            except Exception as e:
                st.error(f"Error al procesar el archivo de audio: {str(e)}")
                st.info("Asegúrate de que FFmpeg esté instalado correctamente en tu sistema.")
                
# Convertidor de hojas de cálculo
def spreadsheet_converter():
    st.subheader("Convertidor de Hojas de Cálculo")
    
    st.info("Actualmente soportamos las siguientes conversiones: CSV a Excel, Excel a CSV, Excel/CSV a PDF")
    
    conversion_type = st.selectbox(
        "Selecciona el tipo de conversión:",
        ["CSV a Excel", "Excel a CSV", "Excel/CSV a PDF"]
    )
    
    if conversion_type == "CSV a Excel":
        csv_to_excel()
    elif conversion_type == "Excel a CSV":
        excel_to_csv()
    elif conversion_type == "Excel/CSV a PDF":
        spreadsheet_to_pdf()

# Funciones para conversión de hojas de cálculo

def csv_to_excel():
    uploaded_file = st.file_uploader("📊 Sube un archivo CSV", type=["csv"])
    
    if uploaded_file:
        try:
            # Leer el CSV
            df = pd.read_csv(uploaded_file)
            
            # Mostrar vista previa
            st.write("Vista previa de los datos:")
            st.dataframe(df.head())
            
            # Opciones de conversión
            sheet_name = st.text_input("Nombre de la hoja de Excel:", "Hoja1")
            
            # Botón de conversión
            if st.button("Convertir a Excel"):
                # Crear un archivo Excel en memoria
                buffer = BytesIO()
                with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                
                buffer.seek(0)
                
                st.success("✅ CSV convertido a Excel")
                st.download_button(
                    label="📥 Descargar Excel",
                    data=buffer,
                    file_name="datos_convertidos.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
                
        except Exception as e:
            st.error(f"Error al convertir el archivo: {str(e)}")

def excel_to_csv():
    uploaded_file = st.file_uploader("📊 Sube un archivo Excel", type=["xlsx", "xls"])
    
    if uploaded_file:
        try:
            # Leer el archivo Excel
            xls = pd.ExcelFile(uploaded_file)
            sheet_names = xls.sheet_names
            
            # Seleccionar una hoja si hay varias
            selected_sheet = st.selectbox("Selecciona la hoja a convertir:", sheet_names)
            
            # Leer la hoja seleccionada
            df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
            
            # Mostrar vista previa
            st.write("Vista previa de los datos:")
            st.dataframe(df.head())
            
            # Opciones de conversión
            separator = st.selectbox("Separador para CSV:", [",", ";", "Tab"])
            if separator == "Tab":
                separator = "\t"
            
            # Opciones de codificación
            encoding = st.selectbox("Codificación:", ["utf-8", "latin-1", "ascii"])
            
            # Botón de conversión
            if st.button("Convertir a CSV"):
                # Crear un archivo CSV en memoria
                buffer = BytesIO()
                df.to_csv(buffer, sep=separator, index=False, encoding=encoding)
                buffer.seek(0)
                
                st.success("✅ Excel convertido a CSV")
                st.download_button(
                    label="📥 Descargar CSV",
                    data=buffer,
                    file_name="datos_convertidos.csv",
                    mime="text/csv"
                )
                
        except Exception as e:
            st.error(f"Error al convertir el archivo: {str(e)}")

def spreadsheet_to_pdf():
    st.write("Convierte Excel o CSV a un archivo PDF")
    
    file_type = st.radio("Selecciona el tipo de archivo a convertir:", ["Excel", "CSV"])
    
    if file_type == "Excel":
        uploaded_file = st.file_uploader("📊 Sube un archivo Excel", type=["xlsx", "xls"])
        if uploaded_file:
            try:
                # Leer el archivo Excel
                xls = pd.ExcelFile(uploaded_file)
                sheet_names = xls.sheet_names
                
                # Seleccionar una hoja si hay varias
                selected_sheet = st.selectbox("Selecciona la hoja a convertir:", sheet_names)
                
                # Leer la hoja seleccionada
                df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
                
                # Mostrar vista previa
                st.write("Vista previa de los datos:")
                st.dataframe(df.head())
                
                process_dataframe_to_pdf(df)
                
            except Exception as e:
                st.error(f"Error al procesar el archivo Excel: {str(e)}")
    
    else:  # CSV
        uploaded_file = st.file_uploader("📊 Sube un archivo CSV", type=["csv"])
        if uploaded_file:
            try:
                # Opciones para leer CSV
                separator = st.selectbox("Separador utilizado en el CSV:", [",", ";", "Tab"])
                if separator == "Tab":
                    separator = "\t"
                
                encoding = st.selectbox("Codificación del archivo:", ["utf-8", "latin-1", "ascii"])
                
                # Leer el CSV
                df = pd.read_csv(uploaded_file, sep=separator, encoding=encoding)
                
                # Mostrar vista previa
                st.write("Vista previa de los datos:")
                st.dataframe(df.head())
                
                process_dataframe_to_pdf(df)
                
            except Exception as e:
                st.error(f"Error al procesar el archivo CSV: {str(e)}")

def process_dataframe_to_pdf(df):
    # Opciones para la conversión a PDF
    orientation = st.selectbox("Orientación de la página:", ["Horizontal", "Vertical"])
    page_size = st.selectbox("Tamaño de página:", ["A4", "Letter", "Legal"])
    
    # Conversión a PDF
    if st.button("Convertir a PDF"):
        try:
            # Crear un archivo PDF temporal
            temp_dir = tempfile.TemporaryDirectory()
            pdf_path = f"{temp_dir.name}/spreadsheet.pdf"
            
            # Configurar el documento PDF
            if page_size == "A4":
                from reportlab.lib.pagesizes import A4
                pagesize = A4
            elif page_size == "Letter":
                from reportlab.lib.pagesizes import letter
                pagesize = letter
            else:  # Legal
                from reportlab.lib.pagesizes import legal
                pagesize = legal
            
            # Ajustar orientación si es necesario
            if orientation == "Horizontal":
                pagesize = pagesize[1], pagesize[0]  # Intercambiar ancho y alto
            
            # Crear el PDF
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=pagesize,
                rightMargin=30,
                leftMargin=30,
                topMargin=30,
                bottomMargin=30
            )
            
            # Preparar los datos para la tabla
            data = [df.columns.tolist()] + df.values.tolist()
            
            # Estilo de la tabla
            table = Table(data)
            style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ])
            table.setStyle(style)
            
            # Construir y guardar el PDF
            elements = []
            elements.append(table)
            doc.build(elements)
            
            # Leer el PDF generado
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            
            st.success("✅ Convertido a PDF correctamente")
            st.download_button(
                label="📥 Descargar PDF",
                data=pdf_bytes,
                file_name="datos_convertidos.pdf",
                mime="application/pdf"
            )
            
            # Limpiar archivos temporales
            temp_dir.cleanup()
        
        except Exception as e:
            st.error(f"Error al generar el PDF: {str(e)}")
# Ejecutar la función seleccionada
if option == "Generador de Código QR":
    qr_generator()
elif option == "Recortar PDF":
    pdf_cutter()
elif option == "Convertidor de Texto a Voz":
    text_to_speech_converter()
elif option == "Convertidor de Archivos":
    file_converter()

# Mostrar información de la aplicación en el pie de página
st.sidebar.markdown("---")
st.sidebar.info(
    """
    **Aplicación Multifuncional v1.0**
    
    Esta aplicación cuenta con múltiples herramientas útiles para el trabajo diario.
    
    Desarrollada con Streamlit.
    """
)
